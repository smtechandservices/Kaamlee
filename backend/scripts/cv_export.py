import io

from django.template.loader import render_to_string
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt


def render_cv_html(content, template):
    return render_to_string('cv/resume.html', {'cv': content, 'template': template})


def render_cv_pdf(content, template):
    html = render_cv_html(content, template)
    buffer = io.BytesIO()
    result = pisa.CreatePDF(html, dest=buffer)
    if result.err:
        return None
    return buffer.getvalue()


def render_cv_docx(content):
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)

    if content.get('name'):
        doc.add_heading(content['name'], level=0)
    if content.get('role'):
        doc.add_paragraph(content['role'])

    contact_bits = [c.get('value', '') for c in (content.get('contacts') or [])]
    contact_bits += [f"{l.get('label', '')}: {l.get('url', '')}" for l in (content.get('links') or [])]
    if contact_bits:
        doc.add_paragraph(' | '.join(contact_bits))

    if content.get('summary'):
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(content['summary'])

    if content.get('skills'):
        doc.add_heading('Skills', level=1)
        for group in content['skills']:
            items = ', '.join(group.get('items') or [])
            category = group.get('category')
            p = doc.add_paragraph()
            if category:
                p.add_run(f"{category}: ").bold = True
            p.add_run(items)

    if content.get('experience'):
        doc.add_heading('Experience', level=1)
        for exp in content['experience']:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('role', '')} — {exp.get('company', '')}").bold = True
            sub = ' | '.join(filter(None, [exp.get('period'), exp.get('location')]))
            if sub:
                doc.add_paragraph(sub)
            for bullet in exp.get('bullets') or []:
                doc.add_paragraph(bullet, style='List Bullet')

    if content.get('projects'):
        doc.add_heading('Projects', level=1)
        for proj in content['projects']:
            p = doc.add_paragraph()
            title = proj.get('name', '')
            if proj.get('tech'):
                title += f" ({', '.join(proj['tech'])})"
            p.add_run(title).bold = True
            if proj.get('description'):
                doc.add_paragraph(proj['description'])
            for bullet in proj.get('bullets') or []:
                doc.add_paragraph(bullet, style='List Bullet')

    if content.get('education'):
        doc.add_heading('Education', level=1)
        for edu in content['education']:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')} — {edu.get('institution', '')}").bold = True
            sub = ' | '.join(filter(None, [edu.get('period'), edu.get('location')]))
            if sub:
                doc.add_paragraph(sub)

    if content.get('certifications'):
        doc.add_heading('Certifications', level=1)
        for cert in content['certifications']:
            line = cert.get('name', '')
            if cert.get('issuer'):
                line += f" — {cert['issuer']}"
            if cert.get('date'):
                line += f" ({cert['date']})"
            doc.add_paragraph(line)

    if content.get('achievements'):
        doc.add_heading('Achievements', level=1)
        for achievement in content['achievements']:
            doc.add_paragraph(achievement, style='List Bullet')

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
